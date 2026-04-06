"""
文档解析器 - 支持多种文档格式
"""

import logging
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

from .chunking import ChunkProcessor, Chunk


class BaseParser(ABC):
    """解析器基类"""

    @abstractmethod
    def can_parse(self, file_path: Path) -> bool:
        """检查是否支持此文件"""
        pass

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """解析文件并返回文本内容"""
        pass


class TextParser(BaseParser):
    """文本文件解析器"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".txt", ".md", ".markdown"]

    def parse(self, file_path: Path) -> str:
        """解析文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            if file_path.suffix.lower() in [".md", ".markdown"]:
                content = self._strip_frontmatter(content)

            return content
        except Exception as e:
            self.logger.error(f"读取文本文件失败: {e}")
            return ""

    def _strip_frontmatter(self, text: str) -> str:
        """移除Markdown YAML frontmatter"""
        import re

        pattern = r"^---\s*\n.*?\n---\s*\n?"
        return re.sub(pattern, "", text, flags=re.DOTALL)


class PDFParser(BaseParser):
    """PDF文件解析器"""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def parse(self, file_path: Path) -> str:
        """解析PDF文件"""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            self.logger.warning("PyMuPDF 未安装，尝试使用 PyPDF2")
            return self._parse_with_pypdf(file_path)
        except Exception as e:
            self.logger.error(f"PDF解析失败: {e}")
            return ""

    def _parse_with_pypdf(self, file_path: Path) -> str:
        """使用 PyPDF2 解析"""
        try:
            import PyPDF2

            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text
        except Exception as e:
            self.logger.error(f"PyPDF2解析失败: {e}")
            return ""


class DocxParser(BaseParser):
    """Word文档解析器"""

    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".doc"]

    def parse(self, file_path: Path) -> str:
        """解析Word文档"""
        try:
            import docx

            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join(cell.text for cell in row.cells) + "\n"
            return text
        except ImportError:
            self.logger.error("python-docx 未安装")
            return ""
        except Exception as e:
            self.logger.error(f"Word文档解析失败: {e}")
            return ""


class DocumentParser:
    """文档解析器管理器"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.parsers: List[BaseParser] = [
            TextParser(),
            PDFParser(),
            DocxParser(),
        ]

    def get_parser(self, file_path: Path) -> Optional[BaseParser]:
        """获取适合的解析器"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None

    def parse(self, file_path: Path) -> str:
        """解析文档"""
        parser = self.get_parser(file_path)
        if parser:
            return parser.parse(file_path)
        self.logger.warning(f"不支持的文件格式: {file_path}")
        return ""


class DocumentProcessor:
    """文档处理器 - 整合解析和分块"""

    def __init__(self, chunk_size: int = 600, overlap: int = 150):
        """
        初始化文档处理器

        Args:
            chunk_size: 分块大小
            overlap: 重叠大小
        """
        self.logger = logging.getLogger(__name__)
        self.parser = DocumentParser()
        self.chunker = ChunkProcessor(chunk_size, overlap)

    def process_file(self, file_path: Path) -> List[Chunk]:
        """
        处理单个文件

        Args:
            file_path: 文件路径

        Returns:
            Chunk 列表
        """
        if not file_path.exists():
            self.logger.error(f"文件不存在: {file_path}")
            return []

        # 解析文档
        text = self.parser.parse(file_path)
        if not text:
            self.logger.warning(f"文档内容为空: {file_path}")
            return []

        # 分块
        chunks = self.chunker.chunk_text(text, source=file_path.name)
        self.logger.info(f"文档处理完成: {file_path}, 生成 {len(chunks)} 个块")

        return chunks

    def process_files(self, file_paths: List[Path]) -> List[Chunk]:
        """
        批量处理文件

        Args:
            file_paths: 文件路径列表

        Returns:
            所有文件的 Chunk 列表
        """
        all_chunks = []
        for path in file_paths:
            chunks = self.process_file(path)
            all_chunks.extend(chunks)
        return all_chunks
